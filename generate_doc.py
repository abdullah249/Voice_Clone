#!/usr/bin/env python3
"""
Generate a comprehensive project documentation .docx file
for the Real-Time Voice Cloning Chatbot project.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import io

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a styled table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    
    return table

def add_code_block(doc, code, language=""):
    """Add a code block with monospace font and grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_diagram_box(doc, title, lines, width=5.5):
    """Create an ASCII-art style diagram rendered as a styled code block."""
    box_lines = []
    max_len = max(len(l) for l in [title] + lines) + 4
    box_lines.append("┌" + "─" * (max_len) + "┐")
    box_lines.append("│  " + title.center(max_len - 4) + "  │")
    box_lines.append("├" + "─" * (max_len) + "┤")
    for l in lines:
        box_lines.append("│  " + l.ljust(max_len - 4) + "  │")
    box_lines.append("└" + "─" * (max_len) + "┘")
    add_code_block(doc, "\n".join(box_lines))


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# ── Customize styles ─────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    style = doc.styles[f'Heading {level}']
    style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    style.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Real-Time Voice Cloning Chatbot")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Project Documentation")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Conversational Agent with Voice Cloning Technology")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph("")

# Project meta table on title page
info_data = [
    ["Project Name", "Real-Time Voice Cloning Chatbot"],
    ["Version", "1.0"],
    ["Date", "February 27, 2026"],
    ["Author", "Abdullah"],
    ["Repository", "https://github.com/abdullah249/Voice_Clone"],
    ["License", "Apache 2.0 / MIT (components)"],
]
table = doc.add_table(rows=len(info_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(table.rows[i].cells[0], "1F4E79")
    table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Introduction",
    "   1.1 Project Overview",
    "   1.2 Objectives",
    "   1.3 Scope",
    "2. System Architecture",
    "   2.1 High-Level Architecture Diagram",
    "   2.2 Component Overview",
    "   2.3 Data Flow Diagram",
    "   2.4 Technology Stack",
    "3. Module Descriptions",
    "   3.1 realtime_voice_chat.py — CLI Voice Chat",
    "   3.2 simple_voice_chat.py — Production Server",
    "   3.3 voice_chat_ui.py — Gradio Web UI",
    "   3.4 voice_admin.py — Admin Panel",
    "   3.5 voice_agent_ari.py — Asterisk Phone Agent",
    "4. Voice Cloning Pipeline",
    "   4.1 Voice Enrollment Process",
    "   4.2 Real-Time Inference Pipeline",
    "   4.3 Pipeline Sequence Diagram",
    "5. API Reference",
    "   5.1 Flask REST API (Port 5050)",
    "   5.2 Gradio Web UI (Port 7860)",
    "   5.3 Admin Panel (Port 9090)",
    "6. Voice Library Management",
    "   6.1 Storage Structure",
    "   6.2 Voice Metadata Schema",
    "7. Deployment Guide",
    "   7.1 Server Requirements",
    "   7.2 Installation Steps",
    "   7.3 Deployment Architecture Diagram",
    "   7.4 Production Deployment (systemd)",
    "   7.5 Reverse Proxy & SSL",
    "8. Configuration & Environment Variables",
    "9. Hardware & GPU Layout",
    "10. Security Considerations",
    "11. Troubleshooting Guide",
    "12. Future Enhancements",
    "13. References & Licenses",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith("   "):
        p.runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Project Overview", level=2)
doc.add_paragraph(
    "The Real-Time Voice Cloning Chatbot is an AI-powered conversational system that enables users "
    "to interact with a large language model (LLM) through natural voice, while the AI responds using "
    "a cloned version of any target voice. The system combines state-of-the-art speech-to-text (STT), "
    "text-to-speech (TTS) with voice cloning, and LLM inference into a unified real-time pipeline."
)
doc.add_paragraph(
    "The project supports multiple interfaces: a web-based Gradio UI for browser access, a command-line "
    "interface for developer use, a Flask REST API for programmatic integration, an admin panel for "
    "managing cloned voices, and an Asterisk ARI agent for telephony (phone call) integration."
)

doc.add_heading("1.2 Objectives", level=2)
objectives = [
    "Enable real-time voice conversation with an AI that speaks in any cloned voice.",
    "Provide a user-friendly web interface for voice cloning and chat interactions.",
    "Support multiple LLM backends (Groq, OpenAI, Ollama) for flexible deployment.",
    "Deliver a production-ready system with systemd integration, auto-restart, and monitoring.",
    "Integrate with Asterisk PBX for phone-based voice agent capabilities.",
    "Maintain a voice library allowing multiple cloned voices with easy switching.",
    "Optimize for both GPU (real-time) and CPU (fallback) inference.",
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("1.3 Scope", level=2)
doc.add_paragraph(
    "This document covers the full system: architecture, module-level descriptions, API reference, "
    "voice cloning pipeline details, deployment procedures, configuration, security, and troubleshooting. "
    "It is intended for developers, system administrators, and stakeholders evaluating or deploying the system."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. System Architecture", level=1)

doc.add_heading("2.1 High-Level Architecture Diagram", level=2)
doc.add_paragraph(
    "The following diagram illustrates the high-level architecture of the system, showing how "
    "users interact through different interfaces and how the internal components communicate."
)

arch_diagram = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                        EXTERNAL INTERFACES                                  │
│                                                                             │
│   ┌──────────────┐   ┌──────────────┐   ┌──────────────┐   ┌────────────┐  │
│   │  Web Browser  │   │  Phone/SIP   │   │  Admin Panel │   │    CLI     │  │
│   │  (Port 7860)  │   │  (Asterisk)  │   │  (Port 9090) │   │  Terminal  │  │
│   └──────┬───────┘   └──────┬───────┘   └──────┬───────┘   └─────┬──────┘  │
│          │                  │                   │                  │         │
└──────────┼──────────────────┼───────────────────┼──────────────────┼─────────┘
           │                  │                   │                  │
           ▼                  ▼                   ▼                  ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                      APPLICATION LAYER                                      │
│                                                                             │
│   ┌──────────────┐   ┌──────────────┐   ┌──────────────┐   ┌────────────┐  │
│   │  Gradio UI   │   │  ARI Agent   │   │ Flask Admin  │   │  CLI Chat  │  │
│   │ voice_chat   │   │ voice_agent  │   │ voice_admin  │   │ realtime_  │  │
│   │  _ui.py      │   │  _ari.py     │   │  .py         │   │ voice_chat │  │
│   └──────┬───────┘   └──────┬───────┘   └──────┬───────┘   └─────┬──────┘  │
│          │                  │                   │                  │         │
│          └──────────┬───────┘                   │                  │         │
│                     ▼                           │                  │         │
│   ┌────────────────────────────────┐            │                  │         │
│   │  Flask REST API (Port 5050)    │◄───────────┘                  │         │
│   │  /api/stt  /api/llm  /api/tts │                                │         │
│   │  /api/pipeline  /api/voices    │                                │         │
│   └──────────────┬─────────────────┘                               │         │
│                  │                                                  │         │
└──────────────────┼──────────────────────────────────────────────────┼─────────┘
                   │                                                  │
                   ▼                                                  ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                          AI / ML ENGINE LAYER                               │
│                                                                             │
│   ┌────────────────┐   ┌────────────────────┐   ┌───────────────────────┐   │
│   │  Whisper STT   │   │   Qwen3-TTS 1.7B   │   │  Voice Clone Prompts  │   │
│   │  (GPU: cuda:1) │   │   (GPU: cuda:0)    │   │  (voices/ directory)  │   │
│   └────────────────┘   └────────────────────┘   └───────────────────────┘   │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘
                   │
                   ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                       EXTERNAL SERVICES (LLM)                               │
│                                                                             │
│   ┌────────────────┐   ┌────────────────┐   ┌────────────────┐             │
│   │   Groq Cloud   │   │  OpenAI API    │   │  Ollama Local  │             │
│   │   (Default)    │   │  (Optional)    │   │  (Port 11434)  │             │
│   └────────────────┘   └────────────────┘   └────────────────┘             │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, arch_diagram)

doc.add_heading("2.2 Component Overview", level=2)
doc.add_paragraph(
    "The system is composed of several key components organized in a layered architecture:"
)

add_styled_table(doc,
    ["Component", "Description", "Port / Interface"],
    [
        ["Gradio Web UI", "Browser-based interface for voice cloning and chat interaction. Users upload voice samples, create clones, and converse via microphone or text.", "7860 (HTTP)"],
        ["Flask REST API", "Internal JSON/file API exposing STT, LLM, TTS, and pipeline endpoints. Used by the ARI agent and admin panel.", "5050 (HTTP, local only)"],
        ["Admin Panel", "Password-protected web dashboard for managing cloned voices — activate, rename, delete, preview voices.", "9090 (HTTP)"],
        ["ARI Voice Agent", "Asterisk Integration via WebSocket. Handles incoming phone calls with real-time voice conversation using the clone pipeline.", "8088 (ARI WS)"],
        ["CLI Voice Chat", "Command-line chatbot for terminal use. Directly records microphone, runs pipeline, and plays audio.", "N/A (local)"],
        ["Whisper STT", "OpenAI Whisper model for speech-to-text transcription. Deployed on a dedicated GPU.", "In-process"],
        ["Qwen3-TTS", "Alibaba Qwen3-TTS model for text-to-speech synthesis with voice cloning capability.", "In-process"],
        ["Voice Library", "Filesystem-based storage of cloned voice prompts (.pt files) with JSON metadata.", "voices/ directory"],
    ],
    col_widths=[1.5, 3.2, 1.5]
)

doc.add_paragraph("")

doc.add_heading("2.3 Data Flow Diagram", level=2)
doc.add_paragraph(
    "The following diagram shows the data flow for a single conversation turn, from user speech input to AI voice response output."
)

flow_diagram = """
┌─────────────────────────────────────────────────────────────────────────┐
│                    CONVERSATION TURN DATA FLOW                          │
└─────────────────────────────────────────────────────────────────────────┘

  ┌──────────┐     Audio (WAV)      ┌──────────────┐     Text
  │          │ ──────────────────►   │              │ ──────────────►
  │   User   │                      │  Whisper STT │      │
  │  (Speak) │                      │  (GPU cuda:1)│      │
  └──────────┘                      └──────────────┘      │
                                                          │
                                                          ▼
                                                   ┌──────────────┐
                                                   │              │
                                                   │   LLM API    │   User Text
                                                   │  (Groq /     │──────────
                                                   │   OpenAI /   │          │
                                                   │   Ollama)    │          │
                                                   └──────┬───────┘          │
                                                          │                  │
                                                          │ AI Response Text │
                                                          ▼                  │
  ┌──────────┐     Audio (WAV)      ┌──────────────┐     Text          Conversation
  │          │ ◄──────────────────   │              │ ◄──────────       History
  │   User   │                      │  Qwen3-TTS   │                  (last 8 msgs)
  │  (Listen)│                      │  + Voice Clone│
  └──────────┘                      │  (GPU cuda:0) │
                                    └──────────────┘"""

add_code_block(doc, flow_diagram)

doc.add_heading("2.4 Technology Stack", level=2)

add_styled_table(doc,
    ["Layer", "Technology", "Purpose"],
    [
        ["Speech-to-Text", "OpenAI Whisper (base/tiny/small/medium/large)", "Transcribe user speech to text"],
        ["Text-to-Speech", "Qwen3-TTS 1.7B / 0.6B (Alibaba)", "Synthesize speech with voice cloning"],
        ["Voice Cloning", "Qwen3-TTS ICL + X-vector modes", "Clone voice from 5–15s audio sample"],
        ["LLM (Default)", "Groq Cloud API (Llama 3.1 8B Instant)", "Generate conversational AI responses"],
        ["LLM (Alt)", "OpenAI API / Ollama (local)", "Alternative LLM providers"],
        ["Web UI", "Gradio 4.x", "Browser-based user interface"],
        ["REST API", "Flask", "Internal service communication"],
        ["Telephony", "Asterisk PBX + ARI WebSocket", "Phone call handling"],
        ["Cloud TTS Fallback", "ElevenLabs / Edge TTS", "Fallback TTS for phone pipeline"],
        ["Deep Learning", "PyTorch (CUDA 12.1)", "GPU-accelerated model inference"],
        ["Audio Processing", "SoX, FFmpeg, sounddevice, soundfile", "Audio format conversion & I/O"],
        ["Server OS", "Ubuntu 22.04 / 24.04 LTS", "Production server operating system"],
        ["Process Manager", "systemd", "Production service management"],
        ["Reverse Proxy", "Nginx (optional)", "SSL termination, domain routing"],
    ],
    col_widths=[1.3, 2.8, 2.1]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. MODULE DESCRIPTIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Module Descriptions", level=1)
doc.add_paragraph(
    "This section describes each Python module in the project, its purpose, key classes/functions, and how it integrates with the overall system."
)

# 3.1 realtime_voice_chat.py
doc.add_heading("3.1 realtime_voice_chat.py — CLI Voice Chat", level=2)
doc.add_paragraph(
    "A self-contained command-line voice chatbot. It initializes Whisper STT, Qwen3-TTS, and an LLM "
    "provider, then enters a conversation loop that records user speech, transcribes it, gets an AI "
    "response, synthesizes it in the cloned voice, and plays it back through the speakers."
)

doc.add_heading("Key Class: RealtimeVoiceChatBot", level=3)
add_styled_table(doc,
    ["Method", "Description"],
    [
        ["__init__()", "Loads Whisper, Qwen3-TTS models; creates voice clone prompt from reference audio."],
        ["_create_voice_prompt()", "Creates voice clone prompt using ICL mode (with transcript) or X-vector mode."],
        ["speech_to_text(audio)", "Transcribes audio numpy array to text using Whisper."],
        ["get_ai_response(text)", "Sends text to configured LLM (Ollama/OpenAI) and returns AI response."],
        ["text_to_speech_cloned(text)", "Generates speech audio in the cloned voice using Qwen3-TTS."],
        ["record_audio()", "Records from microphone with voice activity detection (VAD) and silence detection."],
        ["chat_turn()", "Performs one full conversation turn: record → STT → LLM → TTS → play."],
        ["run()", "Main loop — calls chat_turn() repeatedly until user says 'quit' or 'bye'."],
    ],
    col_widths=[2.2, 4.0]
)

doc.add_paragraph("")
doc.add_paragraph("Command-line arguments:")
add_code_block(doc, """python realtime_voice_chat.py \\
    --clone-audio "voice_sample.wav" \\
    --clone-text "Hello, this is a sample recording." \\
    --llm ollama \\
    --llm-model qwen2.5:7b \\
    --whisper-model base \\
    --tts-model Qwen/Qwen3-TTS-12Hz-1.7B-Base \\
    --device cuda:0""")

# 3.2 simple_voice_chat.py
doc.add_heading("3.2 simple_voice_chat.py — Production Server (Main Entry Point)", level=2)
doc.add_paragraph(
    "The primary production module. It serves the Gradio web UI on port 7860 and a Flask REST API on "
    "port 5050 simultaneously. It pre-loads models at startup, auto-loads saved voice prompts, "
    "manages a voice library, and supports multi-GPU deployment. This is the module deployed as a "
    "systemd service in production."
)

doc.add_heading("Key Features", level=3)
features = [
    "Single-instance lock (prevents duplicate processes via file lock).",
    "Multi-GPU support: TTS on cuda:0 (RTX 4000), Whisper on cuda:1.",
    "Auto-loads saved voice prompt (voice_prompt.pt) on restart — zero downtime voice switching.",
    "Voice library: saves each cloned voice with metadata to voices/ directory.",
    "Integrated Flask API for phone pipeline and admin panel communication.",
    "Groq, OpenAI, and Ollama LLM provider support with UI dropdown.",
    "Audio auto-trimming (up to 15 seconds) for optimal voice cloning quality.",
    "Dynamic max_new_tokens calculation based on text length (12Hz codec).",
    "Per-call conversation history for phone pipeline (keyed by call_id).",
    "ElevenLabs + Edge TTS fallback chain for phone pipeline reliability.",
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("Module Structure Diagram", level=3)

module_diagram = """
┌─────────────────────────── simple_voice_chat.py ──────────────────────────┐
│                                                                           │
│   ┌───────────────────────────────────────────────────────────────┐       │
│   │                    STARTUP / MODEL LOADING                    │       │
│   │  • Single-instance lock (fcntl)                               │       │
│   │  • Load Qwen3-TTS 1.7B → cuda:0 (float16)                    │       │
│   │  • Load Whisper base → cuda:1                                 │       │
│   │  • Auto-load voice_prompt.pt if exists                        │       │
│   └───────────────────────────────────────────────────────────────┘       │
│                                                                           │
│   ┌─────────────────────────┐   ┌─────────────────────────────────┐       │
│   │    GRADIO UI (:7860)    │   │     FLASK API (:5050)           │       │
│   │                         │   │                                 │       │
│   │  • Setup Tab            │   │  GET  /api/health               │       │
│   │    - Model selection    │   │  POST /api/stt                  │       │
│   │    - Voice cloning      │   │  POST /api/llm                  │       │
│   │    - LLM config         │   │  POST /api/tts                  │       │
│   │                         │   │  POST /api/pipeline             │       │
│   │  • Chat Tab             │   │  POST /api/end_call             │       │
│   │    - Voice selector     │   │  GET  /api/voices               │       │
│   │    - Mic input          │   │  GET  /api/voices/active        │       │
│   │    - Text input         │   │  POST /api/voices/activate      │       │
│   │    - Audio playback     │   │                                 │       │
│   └─────────────────────────┘   └─────────────────────────────────┘       │
│                                                                           │
│   ┌───────────────────────────────────────────────────────────────┐       │
│   │                    VOICE LIBRARY SYSTEM                       │       │
│   │  • _save_voice_to_library()   • _list_voices()                │       │
│   │  • _activate_voice()          • _get_active_voice_name()      │       │
│   │  • voices/ directory with meta.json + prompt.pt per voice     │       │
│   └───────────────────────────────────────────────────────────────┘       │
│                                                                           │
└───────────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, module_diagram)

# 3.3
doc.add_heading("3.3 voice_chat_ui.py — Gradio Web UI (Standalone)", level=2)
doc.add_paragraph(
    "A standalone Gradio-based web interface with a VoiceChatUI class. It provides a clean two-tab "
    "interface (Setup + Voice Chat) and supports Ollama and OpenAI as LLM providers. This module is "
    "designed for individual/development use, while simple_voice_chat.py is the production version."
)
add_styled_table(doc,
    ["Tab", "Features"],
    [
        ["Setup", "Model loading (TTS model selection, Whisper size), voice clone creation (audio upload + optional transcript), LLM provider/model configuration."],
        ["Voice Chat", "Microphone recording + send, text input + send, audio playback of AI response, transcription display, conversation history, clear history button."],
    ],
    col_widths=[1.2, 5.0]
)

# 3.4
doc.add_heading("3.4 voice_admin.py — Admin Panel", level=2)
doc.add_paragraph(
    "A Flask-based admin dashboard running on port 9090 with password authentication. It communicates "
    "with the Flask REST API (port 5050) to manage cloned voices. Designed for administrators to "
    "monitor and control the voice pipeline without using the Gradio UI."
)

admin_diagram = """
┌────────────────────── voice_admin.py (:9090) ──────────────────────┐
│                                                                     │
│   ┌─────────────────────────────────────────────────────────────┐   │
│   │  Authentication Layer (session-based, password protected)   │   │
│   └─────────────┬───────────────────────────────────────────────┘   │
│                 │                                                    │
│   ┌─────────────▼───────────────────────────────────────────────┐   │
│   │  Dashboard                                                  │   │
│   │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐        │   │
│   │  │ Total Voices │ │ Active Voice │ │ TTS Model    │        │   │
│   │  │    count     │ │    name      │ │   Status     │        │   │
│   │  └──────────────┘ └──────────────┘ └──────────────┘        │   │
│   │                                                             │   │
│   │  Voice Cards:                                               │   │
│   │  ┌──────────────────────────────────────────────┐           │   │
│   │  │ Voice Name        [● ACTIVE]                 │           │   │
│   │  │ Created: 2026-02-19  Model: 1.7B             │           │   │
│   │  │ [Activate] [Sample] [Preview] [Rename] [Del] │           │   │
│   │  └──────────────────────────────────────────────┘           │   │
│   └─────────────────────────────────────────────────────────────┘   │
│                 │                                                    │
│                 │  HTTP requests to Flask API                        │
│                 ▼                                                    │
│   ┌─────────────────────────────────────────────────────────────┐   │
│   │  Flask REST API (127.0.0.1:5050)                            │   │
│   │  /api/voices  /api/voices/activate  /api/voices/rename ...  │   │
│   └─────────────────────────────────────────────────────────────┘   │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, admin_diagram)

# 3.5
doc.add_heading("3.5 voice_agent_ari.py — Asterisk Phone Agent", level=2)
doc.add_paragraph(
    "An asynchronous ARI (Asterisk REST Interface) WebSocket agent that handles incoming phone calls. "
    "It connects to Asterisk via WebSocket, answers calls, records caller speech with silence detection, "
    "sends audio through the Flask pipeline, and plays the AI's cloned-voice response back to the caller."
)

doc.add_heading("Phone Call Flow Diagram", level=3)

phone_diagram = """
┌──────────────────────── Phone Call Flow ─────────────────────────────┐
│                                                                      │
│   Caller                    Asterisk ARI              Voice Pipeline │
│     │                           │                          │         │
│     │── Incoming Call ──────►   │                          │         │
│     │                           │── StasisStart ──►  Agent │         │
│     │   ◄── Answer + Beep ──   │                          │         │
│     │                           │                          │         │
│     │── Speak ──────────────►   │── Record ──────►  Agent  │         │
│     │                           │   (silence det.)         │         │
│     │                           │                          │         │
│     │                           │── RecordingDone ► Agent  │         │
│     │   ◄── Processing Beep ─   │                          │         │
│     │                           │      ┌───────────────────┤         │
│     │                           │      │  Flask Pipeline   │         │
│     │                           │      │  1. Groq STT      │         │
│     │                           │      │  2. Groq LLM      │         │
│     │                           │      │  3. Cloned TTS    │         │
│     │                           │      │  → 8kHz WAV out   │         │
│     │                           │      └──────┬────────────┤         │
│     │   ◄── Play TTS WAV ────   │  ◄──────────┘           │         │
│     │                           │                          │         │
│     │   (PlaybackFinished)      │── Beep + Record ► Agent  │         │
│     │                           │   (next turn)            │         │
│     │                           │                          │         │
│     │── Hang Up ────────────►   │── StasisEnd ────► Cleanup│         │
│     │                           │                          │         │
└──────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, phone_diagram)

doc.add_paragraph("")
doc.add_paragraph("ARI Configuration:")
add_styled_table(doc,
    ["Parameter", "Value", "Description"],
    [
        ["ARI_HOST", "127.0.0.1", "Asterisk host address"],
        ["ARI_PORT", "8088", "ARI HTTP port"],
        ["ARI_USER", "voiceagent", "ARI authentication user"],
        ["ARI_APP", "voiceagent", "Stasis application name"],
        ["MAX_SILENCE_SEC", "2", "Seconds of silence before stopping recording"],
        ["MAX_RECORD_SEC", "10", "Maximum recording duration per turn"],
        ["DTMF_TERMINATE", "#", "DTMF key to end recording early"],
        ["FLASK_API", "http://127.0.0.1:5050", "Pipeline backend URL"],
    ],
    col_widths=[1.5, 2.0, 2.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. VOICE CLONING PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Voice Cloning Pipeline", level=1)

doc.add_heading("4.1 Voice Enrollment Process", level=2)
doc.add_paragraph(
    "Voice enrollment (cloning) is the process of creating a reusable voice prompt from a short "
    "audio sample. The system supports two modes:"
)

add_styled_table(doc,
    ["Mode", "Input Required", "Quality", "Use Case"],
    [
        ["ICL (In-Context Learning)", "Audio sample + text transcript", "High — captures prosody, rhythm, and timbre", "When transcript is available"],
        ["X-Vector Only", "Audio sample only", "Good — captures speaker characteristics via embeddings", "When no transcript is available"],
    ],
    col_widths=[1.5, 1.8, 1.8, 1.5]
)

doc.add_paragraph("")
doc.add_paragraph("Enrollment steps:")
enrollment_steps = [
    "User uploads a 5–15 second audio clip (WAV, MP3, FLAC, or OGG).",
    "Audio is auto-trimmed to 15 seconds if longer (trimmed from start for cleaner speech).",
    "Qwen3-TTS creates a voice clone prompt tensor using the selected mode.",
    "Prompt is saved as voice_prompt.pt (active) and to voices/<name>/prompt.pt (library).",
    "Metadata (name, creation time, model version, source audio path) is saved as meta.json.",
    "Voice is automatically set as the active voice for all pipeline endpoints.",
]
for i, step in enumerate(enrollment_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("4.2 Real-Time Inference Pipeline", level=2)
doc.add_paragraph(
    "Once a voice is enrolled, the inference pipeline processes each conversation turn:"
)

pipeline_diagram = """
┌──────────────────────── Inference Pipeline ──────────────────────────┐
│                                                                      │
│  ┌─────────┐   ┌──────────┐   ┌──────────┐   ┌──────────────────┐    │
│  │  Audio  │──►│ Whisper  │──►│   LLM    │──►│   Qwen3-TTS +    │    │
│  │  Input  │   │   STT    │   │ (Groq /  │   │   Voice Clone    │    │
│  │  (WAV)  │   │          │   │ OpenAI / │   │   Prompt         │    │
│  │         │   │ "Hello"  │   │ Ollama)  │   │                  │    │
│  └─────────┘   └──────────┘   └──────────┘   └────────┬─────────┘    │
│                                                       │              │
│                  Optimization Details:                ▼              │
│  • Whisper model: tiny→large (speed vs accuracy)   ┌──────────┐      │
│  • LLM max_tokens: 60 (short, conversational)      │  Audio   │      │
│  • TTS max_tokens: dynamic (word_count × 15)       │  Output  │      │
│  • Text capped at 30 words for voice output        │  (WAV)   │      │
│  • 12Hz codec: ~12 tokens per second of audio      └──────────┘      │
│                                                                      │
│  Pipeline for Phone (Asterisk):                                      │
│  • Cloud STT (Groq Whisper API) → faster than local                  │
│  • Cloud LLM (Groq) → low latency                                    │
│  • Local TTS (cloned voice) → falls back to ElevenLabs → Edge TTS    │
│  • Output converted to 8kHz 16-bit WAV for Asterisk compatibility    │
│                                                                      │
└──────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, pipeline_diagram)

doc.add_heading("4.3 Pipeline Sequence Diagram (Web UI)", level=3)

seq_diagram = """
  User          Gradio UI       Whisper STT      LLM (Groq)     Qwen3-TTS
   │                │                │                │               │
   │── Record ─────►│                │                │               │
   │                │── Transcribe ─►│                │               │
   │                │                │── "text" ─────►│               │
   │                │                │                │               │
   │                │                │                │── Generate ──►│
   │                │                │                │   Response    │
   │                │                │  ◄── response──│               │
   │                │                │       text     │               │
   │                │  ◄─────────────│                │── Synthesize ►│
   │                │    AI text     │                │   Clone Voice │
   │                │                │                │               │
   │                │   ◄─────────────────────────────────── Audio ───│
   │   ◄── Play ────│                │                │               │
   │    Audio       │                │                │               │"""

add_code_block(doc, seq_diagram)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. API Reference", level=1)

doc.add_heading("5.1 Flask REST API (Port 5050)", level=2)
doc.add_paragraph(
    "The Flask API runs on 127.0.0.1:5050 (internal only) and is consumed by the ARI agent and admin panel."
)

# Health
doc.add_heading("GET /api/health", level=3)
doc.add_paragraph("Returns system health status.")
add_code_block(doc, """Response:
{
  "status": "ok",
  "whisper_loaded": true,
  "tts_loaded": true,
  "voice_cloned": true
}""")

# STT
doc.add_heading("POST /api/stt", level=3)
doc.add_paragraph("Speech-to-text transcription. Accepts multipart WAV file upload.")
add_styled_table(doc,
    ["Parameter", "Type", "Description"],
    [
        ["audio (file)", "multipart/form-data", "WAV audio file to transcribe"],
    ],
    col_widths=[1.5, 1.5, 3.2]
)
add_code_block(doc, """Response: { "text": "Hello, how are you?" }""")

# LLM
doc.add_heading("POST /api/llm", level=3)
doc.add_paragraph("Get LLM response with per-call conversation history.")
add_styled_table(doc,
    ["Parameter", "Type", "Description"],
    [
        ["text", "string (JSON body)", "User message text"],
        ["call_id", "string (optional)", "Unique call identifier for conversation tracking"],
    ],
    col_widths=[1.5, 1.5, 3.2]
)
add_code_block(doc, """Response: { "response": "I'm doing great!", "call_id": "default" }""")

# TTS
doc.add_heading("POST /api/tts", level=3)
doc.add_paragraph("Text-to-speech with the active cloned voice. Returns a WAV file.")
add_styled_table(doc,
    ["Parameter", "Type", "Description"],
    [
        ["text", "string (JSON body)", "Text to synthesize (capped at 30 words)"],
    ],
    col_widths=[1.5, 1.5, 3.2]
)
doc.add_paragraph("Returns: audio/wav file (16kHz)")

# Pipeline
doc.add_heading("POST /api/pipeline", level=3)
doc.add_paragraph(
    "Full audio-in → audio-out pipeline. Accepts WAV, runs STT → LLM → TTS, returns 8kHz WAV "
    "suitable for Asterisk playback. Uses cloud APIs (Groq Whisper + Groq LLM) with local/cloud TTS fallback chain."
)
add_styled_table(doc,
    ["Parameter", "Type", "Description"],
    [
        ["audio (file)", "multipart/form-data", "Input WAV audio file"],
        ["call_id", "form field", "Unique call identifier"],
    ],
    col_widths=[1.5, 1.5, 3.2]
)
doc.add_paragraph("Response headers: X-TTS-Engine (cloned-voice | elevenlabs | edge-tts), X-AI-Response")

# Voices
doc.add_heading("GET /api/voices", level=3)
doc.add_paragraph("List all saved voice clones with metadata.")
add_code_block(doc, """Response:
{
  "voices": [
    {
      "name": "voice_20260219_114246",
      "created": "20260219_114246",
      "model": "Qwen3-TTS-12Hz-1.7B-Base",
      "dir_name": "voice_20260219_114246",
      "is_active": true,
      "has_sample": true
    }
  ]
}""")

doc.add_heading("POST /api/voices/activate", level=3)
doc.add_paragraph("Set a voice as the active voice for all pipeline endpoints.")
add_code_block(doc, """Request:  { "dir_name": "voice_20260219_114246" }
Response: { "status": "ok", "message": "Voice 'voice_20260219_114246' activated" }""")

doc.add_heading("POST /api/end_call", level=3)
doc.add_paragraph("Clean up per-call conversation history when a phone call ends.")
add_code_block(doc, """Request:  { "call_id": "abc123" }
Response: { "status": "ok" }""")

doc.add_heading("5.2 Gradio Web UI (Port 7860)", level=2)
doc.add_paragraph(
    "The Gradio UI is accessible at http://<server-ip>:7860. It provides a browser-based interface "
    "with two tabs: Setup (model loading, voice cloning, LLM configuration) and Chat (voice/text "
    "input, audio output, conversation history). Supports WebSocket for real-time interaction."
)

doc.add_heading("5.3 Admin Panel (Port 9090)", level=2)
doc.add_paragraph(
    "Password-protected Flask web app at http://<server-ip>:9090. Provides a dashboard showing "
    "total voices, active voice, TTS model status, and voice cards with activate/preview/rename/delete actions."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. VOICE LIBRARY MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Voice Library Management", level=1)

doc.add_heading("6.1 Storage Structure", level=2)

storage_diagram = """
voices/
├── .active                          ← Text file: name of currently active voice
├── voice_20260219_114246/
│   ├── meta.json                    ← Voice metadata (name, created, model)
│   ├── prompt.pt                    ← PyTorch tensor (voice clone prompt)
│   └── sample.mp3                   ← Original audio sample (if saved)
├── voice_20260219_120158/
│   ├── meta.json
│   ├── prompt.pt
│   └── sample.wav
├── voice_20260219_124251/
│   ├── meta.json
│   └── prompt.pt
└── ...

voice_prompt.pt                      ← Copy of active voice's prompt.pt
                                       (loaded at startup for instant availability)"""

add_code_block(doc, storage_diagram)

doc.add_heading("6.2 Voice Metadata Schema (meta.json)", level=2)

add_styled_table(doc,
    ["Field", "Type", "Description"],
    [
        ["name", "string", "Human-readable voice name (also used as directory name)"],
        ["created", "string", "Timestamp in YYYYmmdd_HHMMSS format"],
        ["created_epoch", "float", "Unix timestamp of creation"],
        ["source_audio", "string | null", "Filename of saved original audio sample (e.g., 'sample.mp3')"],
        ["model", "string", "TTS model used for cloning (e.g., 'Qwen3-TTS-12Hz-1.7B-Base')"],
    ],
    col_widths=[1.5, 1.2, 3.5]
)

doc.add_paragraph("")
doc.add_paragraph("Example meta.json:")
add_code_block(doc, """{
  "name": "voice_20260219_114246",
  "created": "20260219_114246",
  "created_epoch": 1771494166.69,
  "source_audio": "sample.mp3",
  "model": "Qwen3-TTS-12Hz-1.7B-Base"
}""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. DEPLOYMENT GUIDE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Deployment Guide", level=1)

doc.add_heading("7.1 Server Requirements", level=2)

doc.add_paragraph("Minimum requirements (CPU-only, slower inference):")
add_styled_table(doc,
    ["Resource", "Minimum", "Recommended (GPU)"],
    [
        ["CPU", "2 vCPUs", "4+ vCPUs"],
        ["RAM", "4 GB (8 GB with swap)", "16 GB"],
        ["GPU", "None (CPU fallback)", "NVIDIA 8+ GB VRAM (T4, A10G, RTX 3060+)"],
        ["Disk", "30 GB", "50 GB"],
        ["OS", "Ubuntu 22.04 LTS", "Ubuntu 22.04 / 24.04 LTS"],
        ["Swap", "8 GB (required for CPU)", "Optional"],
    ],
    col_widths=[1.2, 2.0, 3.0]
)

doc.add_heading("7.2 Installation Steps", level=2)

install_steps = [
    ("System Setup", "sudo apt update && sudo apt upgrade -y\nsudo apt install -y python3 python3-pip python3-venv git curl wget ffmpeg sox libsox-dev build-essential libsndfile1 portaudio19-dev"),
    ("Clone Repository", "cd ~ && git clone https://github.com/abdullah249/Voice_Clone.git && cd Voice_Clone"),
    ("Python Environment", "python3 -m venv .venv && source .venv/bin/activate\npip install --upgrade pip setuptools wheel"),
    ("Install PyTorch (GPU)", "pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu121"),
    ("Install Qwen3-TTS", "pip install git+https://github.com/QwenLM/Qwen3-TTS.git"),
    ("Install Dependencies", "pip install openai-whisper sounddevice soundfile numpy librosa gradio requests openai flask"),
    ("Configure API Keys", "cp .env.example .env && nano .env  # Set GROQ_API_KEY"),
    ("Run Application", "set -a; source .env; set +a\npython3 simple_voice_chat.py"),
]
for title, cmd in install_steps:
    doc.add_paragraph(title, style='List Number')
    add_code_block(doc, cmd)

doc.add_heading("7.3 Deployment Architecture Diagram", level=2)

deploy_diagram = """
┌──────────────────────────── Production Server ─────────────────────────────┐
│                                                                            │
│   ┌──── Internet ──────────────────────────────────────────────────────┐   │
│   │  Users → http(s)://your-domain.com                                 │   │
│   └──────────┬─────────────────────────────────────────────────────────┘   │
│              │                                                             │
│              ▼                                                             │
│   ┌──── Nginx (Optional) ──────────────────────────────────────────────┐   │
│   │  • Reverse proxy (port 80/443 → 7860)                              │   │
│   │  • SSL termination (Let's Encrypt)                                 │   │
│   │  • WebSocket upgrade support                                       │   │
│   └──────────┬─────────────────────────────────────────────────────────┘   │
│              │                                                             │
│              ▼                                                             │
│   ┌──── systemd: voice-clone.service ──────────────────────────────────┐   │
│   │                                                                    │   │
│   │  simple_voice_chat.py                                              │   │
│   │  ├── Gradio UI ──── 0.0.0.0:7860 (public)                          │   │
│   │  └── Flask API ──── 127.0.0.1:5050 (internal)                      │   │
│   │                                                                    │   │
│   └────────────────────────────────────────────────────────────────────┘   │
│              │                                                             │
│   ┌──── systemd: voice-admin.service (optional) ──┐                        │
│   │  voice_admin.py ── 0.0.0.0:9090               │                        │
│   └───────────────────────────────────────────────┘                        │
│              │                                                             │
│   ┌──── Asterisk + ARI Agent (optional) ──────────┐                        │
│   │  voice_agent_ari.py ── WebSocket to :8088     │                        │
│   └───────────────────────────────────────────────┘                        │
│                                                                            │
│   ┌──── GPU Resources ────────────────────────────────────────────────┐    │
│   │  cuda:0 (RTX 4000) → Qwen3-TTS 1.7B (float16)                     │    │
│   │  cuda:1 (RTX 4000) → Whisper STT (base)                           │    │
│   │  cuda:2,3 (P4000)  → Available for scaling                        │    │
│   └───────────────────────────────────────────────────────────────────┘    │
│                                                                            │
│   ┌──── External Services ────────────────────────────────────────────┐    │
│   │  Groq API → LLM + Cloud STT         (required)                    │    │
│   │  ElevenLabs → Fallback TTS           (optional)                   │    │
│   │  Ollama → Local LLM (:11434)         (optional)                   │    │
│   └───────────────────────────────────────────────────────────────────┘    │
│                                                                            │
└────────────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, deploy_diagram)

doc.add_heading("7.4 Production Deployment (systemd)", level=2)

doc.add_paragraph("Create a systemd service file:")
add_code_block(doc, """[Unit]
Description=Voice Clone Chatbot (Gradio + Flask)
After=network.target
Wants=network-online.target

[Service]
Type=simple
User=clone
Group=clone
WorkingDirectory=/home/clone/Voice_Clone
EnvironmentFile=/home/clone/Voice_Clone/.env
ExecStart=/home/clone/Voice_Clone/.venv/bin/python3 simple_voice_chat.py
Restart=on-failure
RestartSec=10
StandardOutput=journal
StandardError=journal
LimitNOFILE=65536
TimeoutStartSec=300

[Install]
WantedBy=multi-user.target""")

doc.add_paragraph("Enable and start:")
add_code_block(doc, """sudo systemctl daemon-reload
sudo systemctl enable voice-clone
sudo systemctl start voice-clone
sudo systemctl status voice-clone""")

doc.add_heading("7.5 Reverse Proxy & SSL", level=2)
doc.add_paragraph("Nginx reverse proxy configuration for domain + SSL:")
add_code_block(doc, """server {
    listen 80;
    server_name your-domain.com;

    proxy_read_timeout 300s;
    proxy_send_timeout 300s;
    client_max_body_size 50M;

    location / {
        proxy_pass http://127.0.0.1:7860;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
    }
}""")

doc.add_paragraph("SSL with Let's Encrypt:")
add_code_block(doc, """sudo apt install -y certbot python3-certbot-nginx
sudo certbot --nginx -d your-domain.com""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Configuration & Environment Variables", level=1)

add_styled_table(doc,
    ["Variable", "Required", "Default", "Description"],
    [
        ["GROQ_API_KEY", "Yes", "(none)", "Groq API key for cloud STT and LLM (https://console.groq.com/keys)"],
        ["ELEVENLABS_API_KEY", "No", "(none)", "ElevenLabs API key for fallback TTS in phone pipeline"],
        ["ELEVENLABS_VOICE_ID", "No", "cjVigY5qzO86Huf0OWal", "ElevenLabs voice ID (default: Eric - Smooth)"],
        ["OPENAI_API_KEY", "No", "(none)", "OpenAI API key if using OpenAI as LLM provider"],
        ["VOICE_ADMIN_PASS", "No", "changeme", "Password for the admin panel (port 9090)"],
        ["VOICE_API_BASE", "No", "http://127.0.0.1:5050", "Base URL for Flask API (used by admin panel)"],
        ["VOICE_ADMIN_PORT", "No", "9090", "Port for the admin panel"],
    ],
    col_widths=[1.6, 0.7, 1.5, 2.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. HARDWARE & GPU LAYOUT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Hardware & GPU Layout", level=1)

doc.add_paragraph(
    "The production server is configured with a multi-GPU setup for optimal performance. "
    "The system intelligently distributes workloads across available GPUs."
)

gpu_diagram = """
┌──────────────────────── GPU Memory Layout ─────────────────────────────┐
│                                                                        │
│  ┌────────────────────────────┐  ┌────────────────────────────┐        │
│  │  cuda:0  —  RTX 4000       │  │  cuda:1  —  RTX 4000       │        │
│  │  (Fast, Tensor Cores)      │  │  (Fast, Tensor Cores)      │        │
│  │                            │  │                            │        │
│  │  ┌──────────────────────┐  │  │  ┌──────────────────────┐  │        │
│  │  │  Qwen3-TTS 1.7B      │  │  │  │  Whisper STT (base)  │  │        │
│  │  │  (float16)           │  │  │  │                      │  │        │
│  │  │  ~3.4 GB VRAM        │  │  │  │  ~1 GB VRAM          │  │        │
│  │  └──────────────────────┘  │  │  └──────────────────────┘  │        │
│  └────────────────────────────┘  └────────────────────────────┘        │
│                                                                        │
│  ┌────────────────────────────┐  ┌────────────────────────────┐        │
│  │  cuda:2  —  P4000          │  │  cuda:3  —  P4000          │        │
│  │  (Available for scaling)   │  │  (Available for scaling)   │        │
│  └────────────────────────────┘  └────────────────────────────┘        │
│                                                                        │
│  CPU: 56 cores   RAM: 62 GB   Threads: 24 (torch) + 4 (interop)        │
│                                                                        │
└────────────────────────────────────────────────────────────────────────┘"""

add_code_block(doc, gpu_diagram)

doc.add_paragraph("")
doc.add_paragraph("Performance optimization settings in simple_voice_chat.py:")
add_styled_table(doc,
    ["Setting", "Value", "Purpose"],
    [
        ["torch.set_num_threads(24)", "24 CPU threads", "Parallel CPU operations (optimized for 56-core server)"],
        ["torch.set_num_interop_threads(4)", "4 interop threads", "Cross-operation parallelism"],
        ["TTS dtype", "float16", "Half-precision for faster GPU inference with minimal quality loss"],
        ["TTS attn_implementation", "sdpa", "Scaled dot-product attention (PyTorch native, no Flash Attention dependency)"],
        ["TTS max_tokens", "dynamic (word × 15)", "Adaptive token budget based on text length"],
        ["Text cap", "30 words", "Prevents excessive TTS generation time"],
    ],
    col_widths=[2.2, 1.5, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. SECURITY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Security Considerations", level=1)

security_items = [
    ("API Key Protection", "All API keys (Groq, OpenAI, ElevenLabs) are loaded from environment variables or .env file. Never hardcoded in source. The .env file should have chmod 600 permissions."),
    ("Admin Panel Authentication", "The voice admin panel (port 9090) requires password authentication. Default password 'changeme' should be changed via VOICE_ADMIN_PASS environment variable."),
    ("Internal API Binding", "The Flask REST API (port 5050) binds to 127.0.0.1 only — not accessible from the internet. Only the Gradio UI (port 7860) is publicly exposed."),
    ("Firewall Configuration", "UFW firewall should be configured to allow only ports 22 (SSH), 7860 (Gradio), and optionally 80/443 (Nginx). Port 5050 and 11434 (Ollama) should remain internal."),
    ("Single-Instance Lock", "Process-level file lock (/tmp/voice_chat.lock) prevents accidental duplicate instances that could cause GPU memory conflicts."),
    ("Input Validation", "Audio uploads are auto-trimmed to 15 seconds. Text inputs to TTS are capped at 30 words. Conversation history is limited to 20 messages."),
    ("HTTPS/SSL", "Production deployments should use Nginx with Let's Encrypt SSL certificates for encrypted browser connections."),
]

for title, desc in security_items:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Troubleshooting Guide", level=1)

troubleshooting = [
    ["CUDA out of memory", "Use smaller models: --whisper-model tiny and TTS 0.6B model. Check GPU usage with nvidia-smi.", "High"],
    ["Ollama connection refused", "Ensure Ollama is running: systemctl status ollama or ollama serve &", "Medium"],
    ["Module not found errors", "Activate virtual environment: source .venv/bin/activate. Verify with pip list.", "Medium"],
    ["No audio input (browser)", "Check microphone permissions in browser. Ensure HTTPS or localhost (required for mic access).", "Medium"],
    ["Port 7860 not accessible", "Check firewall: sudo ufw status. Check if process is running: sudo ss -tlnp | grep 7860.", "High"],
    ["Slow inference on CPU", "Expected: TTS takes 30-60s per response on CPU. Use GPU for real-time performance.", "Low"],
    ["First startup very slow", "Normal: downloads TTS model (~3.5 GB) and Whisper model. Check disk space with df -h.", "Low"],
    ["SoX not found (Linux)", "Install via apt: sudo apt install -y sox libsox-dev. Don't use the Windows sox-14.4.2/ folder.", "Medium"],
    ["Voice prompt not loading", "Verify voice_prompt.pt exists and is not corrupted. Re-clone the voice via Gradio UI.", "Medium"],
    ["Pipeline timeout on phone", "Cloud APIs may have latency. Check GROQ_API_KEY. Pipeline has 60s timeout.", "High"],
]

add_styled_table(doc,
    ["Issue", "Solution", "Severity"],
    troubleshooting,
    col_widths=[1.5, 3.7, 1.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Future Enhancements", level=1)

enhancements = [
    ("Streaming TTS", "Implement streaming audio generation to reduce time-to-first-byte for voice responses."),
    ("Multi-language Support", "Extend voice cloning and conversation to support multiple languages beyond English."),
    ("Voice Quality Scoring", "Add automated quality assessment for cloned voices to guide users toward better samples."),
    ("WebRTC Integration", "Replace HTTP-based audio transfer with WebRTC for lower latency real-time communication."),
    ("User Authentication", "Add multi-user support with individual voice libraries and conversation histories."),
    ("Model Fine-tuning", "Support fine-tuning TTS model on specific voices for higher fidelity cloning."),
    ("Horizontal Scaling", "Support multiple worker processes with load balancing for high-traffic deployments."),
    ("Monitoring Dashboard", "Add Prometheus/Grafana metrics for inference latency, GPU utilization, and API throughput."),
]

for title, desc in enhancements:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. REFERENCES & LICENSES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("13. References & Licenses", level=1)

doc.add_heading("Open-Source Components", level=2)

add_styled_table(doc,
    ["Component", "License", "URL"],
    [
        ["Qwen3-TTS", "Apache 2.0", "https://github.com/QwenLM/Qwen3-TTS"],
        ["OpenAI Whisper", "MIT", "https://github.com/openai/whisper"],
        ["Ollama", "MIT", "https://ollama.ai/"],
        ["Gradio", "Apache 2.0", "https://gradio.app/"],
        ["Flask", "BSD-3-Clause", "https://flask.palletsprojects.com/"],
        ["PyTorch", "BSD-3-Clause", "https://pytorch.org/"],
        ["SoX", "GPL / LGPL", "https://sox.sourceforge.net/"],
        ["FFmpeg", "LGPL / GPL", "https://ffmpeg.org/"],
        ["aiohttp", "Apache 2.0", "https://docs.aiohttp.org/"],
    ],
    col_widths=[1.5, 1.2, 3.5]
)

doc.add_heading("Cloud Service APIs", level=2)

add_styled_table(doc,
    ["Service", "Purpose", "URL"],
    [
        ["Groq", "Cloud LLM inference and Whisper STT", "https://console.groq.com/"],
        ["OpenAI", "Alternative LLM provider", "https://platform.openai.com/"],
        ["ElevenLabs", "Fallback cloud TTS", "https://elevenlabs.io/"],
        ["Edge TTS", "Free fallback TTS (Microsoft)", "https://github.com/rany2/edge-tts"],
    ],
    col_widths=[1.5, 2.0, 2.7]
)

doc.add_paragraph("")
doc.add_paragraph("")

# Final centered note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Voice_Cloning_Project_Documentation.docx")
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
